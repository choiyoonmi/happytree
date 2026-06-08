from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(95, 99, 104)


def build_exam_docx(exam, output_path, teacher=False):
    document = Document()
    configure_document(document)
    add_cover(document, exam, teacher)
    document.add_section(WD_SECTION.NEW_PAGE)

    current_part = None
    for question in exam["questions"]:
        if question["part"] != current_part:
            current_part = question["part"]
            add_part_heading(document, current_part)
        add_question(document, question, teacher)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE


def add_cover(document, exam, teacher):
    for _ in range(4):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(exam["title"])
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("교사용 정답 및 해설" if teacher else "학생용 문제지")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = GRAY

    document.add_paragraph()
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"총 {len(exam['questions'])}문항  |  이름: ____________________")

    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(20)
    notice.add_run("지문을 읽고 각 문항의 지시에 따라 답하시오.")


def add_part_heading(document, part):
    names = {
        "PART1": "PART 1. 어법 / 어휘 변형",
        "PART2": "PART 2. 문장 순서 / 삽입",
        "PART3": "PART 3. 제목 / 주제 / 요약",
        "PART4": "PART 4. 서술형",
    }
    heading = document.add_heading(names.get(part, part), level=1)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(8)


def add_question(document, question, teacher):
    heading = document.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    source = question.get("passage_number")
    source_label = f"원문 {source} · " if source else ""
    number = heading.add_run(
        f"{question['number']}. [{source_label}{question['type']}] "
    )
    number.bold = True
    number.font.color.rgb = BLUE
    heading.add_run(question["prompt"])

    passage = document.add_paragraph()
    set_box(passage, "E8EEF5")
    passage.paragraph_format.left_indent = Inches(0.12)
    passage.paragraph_format.right_indent = Inches(0.12)
    passage.add_run(question["passage"])

    for index, option in enumerate(question.get("options", []), start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.add_run(f"{index}. {option}")

    if not teacher:
        answer_line = document.add_paragraph()
        answer_line.paragraph_format.space_after = Pt(12)
        answer_line.add_run("답: ______________________________________________")
    else:
        answer = document.add_paragraph()
        set_box(answer, "FFF2CC")
        label = answer.add_run(f"정답: {question['answer']}\n")
        label.bold = True
        answer.add_run(f"해설: {question['explanation']}")


def set_box(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "D9E2F3")
        borders.append(tag)
    p_pr.append(borders)
